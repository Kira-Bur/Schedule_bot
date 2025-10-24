import os
import re
import tempfile
import requests
import io
from PIL import Image, ImageDraw, ImageFont
from docx import Document
import xml.etree.ElementTree as ET
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ImageProcessor:
    def __init__(self):
        self.FONT_SMALL = self.get_russian_font(10)
        self.FONT_REGULAR = self.get_russian_font(14)
    
    def get_russian_font(self, size=10):
        """Получение шрифта с поддержкой русского языка"""
        try:
            font_paths = [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
                '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf',
                'C:/Windows/Fonts/arial.ttf',
                '/System/Library/Fonts/SFNSDisplay.ttf'
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    return ImageFont.truetype(font_path, size)
            
            try:
                response = requests.get('https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf', timeout=10)
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.ttf') as f:
                        f.write(response.content)
                        font = ImageFont.truetype(f.name, size)
                        os.unlink(f.name)
                        return font
            except:
                pass
                
        except Exception as e:
            logger.error(f"Ошибка при загрузке шрифта: {e}")
        
        return ImageFont.load_default()
    
    def get_text_dimensions(self, text, font):
        """Получение размеров текста"""
        try:
            bbox = font.getbbox(text)
            return bbox[2] - bbox[0], bbox[3] - bbox[1]
        except:
            try:
                return font.getsize(text)
            except:
                return len(text) * 6, 12
    
    def calculate_table_width(self, table_data, font):
        """Рассчитывает оптимальную ширину таблицы"""
        if not table_data or not table_data[0]:
            return 0
        
        col_widths = [0] * len(table_data[0])
        for row in table_data:
            for i, cell in enumerate(row):
                if i < len(col_widths):
                    # Находим самое длинное слово в ячейке
                    words = cell.split()
                    max_word_width = 0
                    for word in words:
                        text_width, _ = self.get_text_dimensions(word, font)
                        max_word_width = max(max_word_width, text_width)
                    col_widths[i] = max(col_widths[i], min(max_word_width + 20, 200))
        
        return sum(col_widths)
    
    def docx_to_image(self, docx_path):
        """Конвертация DOCX файла в изображение"""
        try:
            doc = Document(docx_path)
            
            # Фильтруем служебный текст
            skip_patterns = [
                r'УТВЕРЖДАЮ',
                r'Заместитель директора',
                r'Зам. директора по УМР',
                r'Заведующая отделением',
                r'Диспетчер __________________Миронова Е.В',
                r'Расписание на сайте',
                r'«\d+»\s+\w+\s+\d{4}\s+года',
                r'\d{1,2}\s+\w+\s+\d{4}\s+года',
                r'Четверг.*неделя',
                r'Понедельник.*неделя',
                r'Вторник.*неделя',
                r'Среда.*неделя',
                r'Пятница.*неделя',
                r'Суббота.*неделя'
            ] 
            
            def process_time_in_text(text):
                """Обработка времени только для текста вне таблиц"""
                time_pattern = r'\b\d{1,2}[:.\-]\d{2}\b'
                time_matches = re.findall(time_pattern, text)
                
                if len(time_matches) > 2:
                    first_time = time_matches[0]
                    last_time = time_matches[-1]
                    
                    def replace_time(match):
                        time_str = match.group()
                        if time_str == first_time or time_str == last_time:
                            return time_str
                        return ""
                    
                    processed_text = re.sub(time_pattern, replace_time, text)
                    processed_text = re.sub(r'\s+', ' ', processed_text).strip()
                    return processed_text
                return text
            
            # Собираем контент
            useful_content = []
            tables_data = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and not any(re.search(pattern, text, re.IGNORECASE) for pattern in skip_patterns):
                    processed_text = process_time_in_text(text)
                    useful_content.append(processed_text)
            
            # Обрабатываем таблицы
            max_table_width = 0
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_text = re.sub(r'\s+', ' ', cell_text)
                        row_data.append(cell_text)
                    if any(cell.strip() for cell in row_data):
                        table_data.append(row_data)
                
                if table_data and len(table_data) > 1:
                    tables_data.append(table_data)
                    table_width = self.calculate_table_width(table_data, self.FONT_SMALL)
                    max_table_width = max(max_table_width, table_width)
            
            if not useful_content and not tables_data:
                return None
            
            # Настройки отображения
            margin = 35
            line_height = 30
            table_spacing = 25
            
            # Определяем ширину изображения на основе таблицы
            if max_table_width > 0:
                image_width = min(max_table_width + margin * 2, 1200)
            else:
                image_width = 600
            
            # Рассчитываем высоту
            total_height = margin * 8
            total_height += len(useful_content) * line_height
            
            for table_data in tables_data:
                table_height = 0
                for row in table_data:
                    max_lines = 1
                    for cell in row:
                        if cell:
                            words = cell.split()
                            current_width = 0
                            lines_needed = 1
                            for word in words:
                                word_width, _ = self.get_text_dimensions(word + " ", self.FONT_SMALL)
                                current_width += word_width
                                if current_width > 150:
                                    lines_needed += 1
                                    current_width = word_width
                            max_lines = max(max_lines, lines_needed)
                    table_height += max_lines * 12 + 6
                total_height += table_height + table_spacing
            
            # Создаем изображение
            img = Image.new('RGB', (image_width, total_height), color='white')
            draw = ImageDraw.Draw(img)
            
            y = margin
            
            # Рисуем текст
            for text in useful_content:
                lines = []
                words = text.split()
                current_line = ""
                
                for word in words:
                    test_line = current_line + word + " "
                    text_width, _ = self.get_text_dimensions(test_line, self.FONT_REGULAR)
                    
                    if text_width < image_width - margin * 2:
                        current_line = test_line
                    else:
                        if current_line:
                            lines.append(current_line.strip())
                        current_line = word + " "
                
                if current_line:
                    lines.append(current_line.strip())
                
                for line in lines:
                    draw.text((margin, y), line, fill='black', font=self.FONT_REGULAR)
                    y += line_height
            
            # Рисуем таблицы
            for table_data in tables_data:
                if y > total_height - 50:
                    break
                    
                y += table_spacing
                
                # Рассчитываем ширину колонок
                col_widths = [0] * len(table_data[0])
                for row in table_data:
                    for i, cell in enumerate(row):
                        if i < len(col_widths) and cell:
                            text_width, _ = self.get_text_dimensions(cell, self.FONT_SMALL)
                            col_widths[i] = max(col_widths[i], min(text_width + 10, 250))
                
                # Корректируем ширину колонок
                total_table_width = sum(col_widths)
                if total_table_width > image_width - margin * 2:
                    min_col_width = 60
                    available_width = image_width - margin * 2
                    
                    for i in range(len(col_widths)):
                        col_widths[i] = max(min_col_width, col_widths[i])
                    
                    if sum(col_widths) > available_width:
                        scale_factor = available_width / sum(col_widths)
                        col_widths = [int(w * scale_factor) for w in col_widths]
                    
                    total_table_width = sum(col_widths)
                
                # Центрируем таблицу
                table_x = margin + (image_width - margin * 2 - total_table_width) // 2
                
                # Рассчитываем высоту таблицы
                row_heights = []
                for row in table_data:
                    max_cell_lines = 1
                    for i, cell in enumerate(row):
                        if i < len(col_widths) and cell:
                            lines = []
                            words = cell.split()
                            current_line = ""
                            
                            for word in words:
                                test_line = current_line + word + " "
                                text_width, _ = self.get_text_dimensions(test_line, self.FONT_SMALL)
                                
                                if text_width < col_widths[i] - 4:
                                    current_line = test_line
                                else:
                                    if current_line:
                                        lines.append(current_line.strip())
                                    current_line = word + " "
                            
                            if current_line:
                                lines.append(current_line.strip())
                            
                            max_cell_lines = max(max_cell_lines, len(lines))
                    
                    row_heights.append(max_cell_lines * 12 + 6)
                
                table_height = sum(row_heights)
                
                # Рисуем границы таблицы
                draw.rectangle([table_x, y, table_x + total_table_width, y + table_height], 
                             outline='black', fill='white', width=1)
                
                # Рисуем ячейки и текст
                current_x = table_x
                for col_idx, col_width in enumerate(col_widths):
                    # Вертикальные линии
                    if col_idx > 0:
                        draw.line([current_x, y, current_x, y + table_height], fill='black', width=1)
                    
                    current_y = y
                    for row_idx, row in enumerate(table_data):
                        # Горизонтальные линии
                        if row_idx > 0:
                            draw.line([current_x, current_y, current_x + col_width, current_y], 
                                     fill='black', width=1)
                        
                        # Текст ячейки
                        if col_idx < len(row):
                            cell_text = row[col_idx]
                            if cell_text:
                                lines = []
                                words = cell_text.split()
                                current_line = ""
                                
                                for word in words:
                                    test_line = current_line + word + " "
                                    text_width, _ = self.get_text_dimensions(test_line, self.FONT_SMALL)
                                    
                                    if text_width < col_width - 4:
                                        current_line = test_line
                                    else:
                                        if current_line:
                                            lines.append(current_line.strip())
                                        current_line = word + " "
                                
                                if current_line:
                                    lines.append(current_line.strip())
                                
                                for line_idx, line_text in enumerate(lines):
                                    if line_idx < 5:
                                        text_y = current_y + 3 + line_idx * 12
                                        draw.text((current_x + 2, text_y), line_text, 
                                                 fill='black', font=self.FONT_SMALL)
                        
                        current_y += row_heights[row_idx]
                    
                    current_x += col_width
                
                y += table_height
            
            # Обрезаем изображение
            img = img.crop((0, 0, image_width, min(y + margin, total_height)))
            
            return img
            
        except Exception as e:
            logger.error(f"Ошибка при конвертации DOCX: {e}")
            return None
    
    def xml_to_image(self, xml_path):
        """Конвертация XML файла в изображение"""
        try:
            tree = ET.parse(xml_path)
            root = tree.getroot()
            
            def process_time_in_text(text):
                """Оставляет только первое и последнее время в строке"""
                time_pattern = r'\b\d{1,2}[:.\-]\d{2}\b'
                time_matches = re.findall(time_pattern, text)
                
                if len(time_matches) > 2:
                    first_time = time_matches[0]
                    last_time = time_matches[-1]
                    
                    def replace_time(match):
                        time_str = match.group()
                        if time_str == first_time or time_str == last_time:
                            return time_str
                        return ""
                    
                    processed_text = re.sub(time_pattern, replace_time, text)
                    processed_text = re.sub(r'\s+', ' ', processed_text).strip()
                    return processed_text
                
                return text
            
            tables_data = []
            
            for table_elem in root.findall('.//table') + root.findall('.//Table'):
                table_data = []
                for row_elem in table_elem.findall('.//row') + table_elem.findall('.//tr'):
                    row_data = []
                    for cell_elem in row_elem.findall('.//cell') + row_elem.findall('.//td'):
                        cell_text = cell_elem.text.strip() if cell_elem.text else ''
                        cell_text = re.sub(r'\s+', ' ', cell_text)
                        processed_text = process_time_in_text(cell_text)
                        row_data.append(processed_text)
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    tables_data.append(table_data)
            
            font = self.FONT_SMALL
            table_spacing = 20
            margin = 10
            max_width = 1600
            
            total_height = margin * 2
            for table in tables_data:
                total_height += len(table) * 20 + table_spacing
            
            img = Image.new('RGB', (max_width, total_height), color='white')
            draw = ImageDraw.Draw(img)
            
            y = margin
            
            for table_data in tables_data:
                y = self.draw_compact_table(draw, table_data, margin, y, font, max_width - margin * 2)
                y += table_spacing
            
            return img
            
        except Exception as e:
            logger.error(f"Ошибка при конвертации XML: {e}")
            return None
    
    def draw_compact_table(self, draw, table_data, x, y, font, max_width=1600):
        """Рисование улучшенной компактной таблицы"""
        if not table_data or not table_data[0]:
            return y
        
        # Определяем оптимальную ширину для каждой колонки
        col_widths = [0] * len(table_data[0])
        for row in table_data:
            for i, cell in enumerate(row):
                if i < len(col_widths):
                    words = cell.split()
                    max_word_width = 0
                    for word in words:
                        text_width, _ = self.get_text_dimensions(word, font)
                        max_word_width = max(max_word_width, text_width)
                    col_widths[i] = max(col_widths[i], min(max_word_width + 10, 150))
        
        # Корректируем ширину колонок
        total_width = sum(col_widths)
        if total_width > max_width:
            scale_factor = max_width / total_width
            col_widths = [int(w * scale_factor) for w in col_widths]
            total_width = sum(col_widths)
        
        row_height = 22
        
        # Рисуем таблицу
        for row_idx, row in enumerate(table_data):
            current_y = y + row_idx * row_height
            
            for col_idx, cell in enumerate(row):
                if col_idx >= len(col_widths):
                    continue
                    
                current_x = x + sum(col_widths[:col_idx])
                cell_width = col_widths[col_idx]
                
                # Рисуем границу ячейки
                draw.rectangle([current_x, current_y, current_x + cell_width, current_y + row_height], 
                              outline='black', fill='white', width=1)
                
                if cell.strip():
                    # Разбиваем текст на строки
                    lines = []
                    words = cell.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + word + " "
                        text_width, _ = self.get_text_dimensions(test_line, font)
                        
                        if text_width < cell_width - 4:
                            current_line = test_line
                        else:
                            if current_line:
                                lines.append(current_line.strip())
                            current_line = word + " "
                    
                    if current_line:
                        lines.append(current_line.strip())
                    
                    # Рисуем текст (максимум 2 строки)
                    for line_idx, line_text in enumerate(lines[:2]):
                        text_x = current_x + 2
                        text_y = current_y + 3 + line_idx * 9
                        draw.text((text_x, text_y), line_text, fill='black', font=font)
        
        return y + len(table_data) * row_height
    
    def convert_to_image(self, file_path):
        """Конвертация файла в изображение"""
        if file_path.lower().endswith('.docx'):
            return self.docx_to_image(file_path)
        elif file_path.lower().endswith('.xml'):
            return self.xml_to_image(file_path)
        return None


